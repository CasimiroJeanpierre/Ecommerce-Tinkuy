# -*- coding: utf-8 -*-
"""Genera el documento Word con las herramientas de accesibilidad para el proyecto Tinkuy."""

import docx
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, url, text):
    """Agrega un hyperlink real (azul y subrayado) a un párrafo de python-docx."""
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_cell_text(cell, text, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


doc = Document()

# Márgenes
section = doc.sections[0]
section.left_margin = Cm(2)
section.right_margin = Cm(2)

# Título
titulo = doc.add_heading('Herramientas de Accesibilidad para Software', level=1)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitulo = doc.add_paragraph()
subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitulo.add_run('Proyecto: E-commerce Tinkuy')
run.italic = True
run.font.size = Pt(12)

doc.add_paragraph()

intro = doc.add_paragraph(
    "Como parte del proyecto, investigamos distintas herramientas que ayudan a "
    "que un sistema de software sea más accesible para personas con discapacidad "
    "visual, motora, auditiva o cognitiva. Para cada una incluimos el nombre y la "
    "versión, el enlace oficial y el tipo de discapacidad al que está orientada."
)

doc.add_paragraph()

# --- Acápite: acciones de accesibilidad ya implementadas (Evidencia ODS 10) ---
heading_ods = doc.add_heading(
    'Acciones de Accesibilidad Implementadas en el Proyecto (ODS 10)',
    level=2,
)

parrafo_ods = doc.add_paragraph(
    "El ODS 10 de la Agenda 2030 busca reducir las desigualdades, lo que "
    "incluye que las personas con discapacidad puedan acceder en igualdad de "
    "condiciones a bienes, servicios y entornos digitales. Pensando en esto, "
    "en Tinkuy ya incorporamos varias prácticas y herramientas de "
    "accesibilidad dentro del código del proyecto:"
)

consideraciones = [
    "Integramos el widget de accesibilidad de UserWay en todas las vistas "
    "públicas (src/Views/components/footer.php). Incluye lector de pantalla, "
    "ajuste de contraste, cambio de tamaño de texto y navegación por teclado, "
    "pensado para usuarios con discapacidad visual, motora y cognitiva.",

    "Las vistas declaran el atributo lang=\"es\" en el HTML, para que los "
    "lectores de pantalla pronuncien el contenido en español correctamente.",

    "Las imágenes del carrusel y de los productos (src/Views/index.php) "
    "tienen atributos alt descriptivos, algo fundamental para quienes "
    "dependen de un lector de pantalla.",

    "En los formularios (por ejemplo src/Views/auth/login.php) las etiquetas "
    "<label> están correctamente asociadas a sus campos mediante el atributo "
    "for, lo que facilita su uso con tecnologías de asistencia.",

    "Agregamos atributos ARIA (aria-label, aria-expanded, aria-labelledby, "
    "role) en menús desplegables, botones y alertas, tanto en el navbar como "
    "en la página de inicio, para mejorar la navegación con lectores de "
    "pantalla.",

    "Definimos estilos de enfoque (focus) visibles para los campos de "
    "formulario en public/css/style.css, útiles para quienes navegan "
    "únicamente con teclado.",

    "El diseño es responsivo gracias al meta viewport y a Bootstrap 5, de "
    "modo que la interfaz se adapta a distintos tamaños de pantalla y "
    "dispositivos de asistencia.",

    "También ajustamos el z-index del menú de navegación, ya que el botón "
    "flotante de accesibilidad tapaba el menú hamburguesa en celulares "
    "(public/css/style.css). Es un detalle pequeño, pero muestra que la "
    "accesibilidad se tuvo en cuenta durante el diseño.",
]

for item in consideraciones:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

subheading_tabla = doc.add_heading(
    'Herramientas de Calidad para Software Inclusivo (Implementadas y Propuestas)',
    level=2,
)

# Datos de las herramientas
herramientas = [
    {
        "nombre": "UserWay Accessibility Widget",
        "version": "Sin versión fija (SaaS) — widget.js servido como última versión vía CDN, cumple WCAG 2.2 AA / ADA / Section 508",
        "url": "https://userway.org",
        "discapacidad": "Visual, motora, cognitiva y auditiva",
        "estado": "Implementada en el proyecto",
    },
    {
        "nombre": "WAVE (Web Accessibility Evaluation Tool)",
        "version": "Extensión de navegador v3.3.1.0 (Chrome Web Store) — WebAIM",
        "url": "https://wave.webaim.org/",
        "discapacidad": "Visual, motora, cognitiva y auditiva",
        "estado": "Recomendada para agregar (QA)",
    },
    {
        "nombre": "Google Lighthouse",
        "version": "v13.0.3 — paquete npm / integrado en Google Chrome DevTools",
        "url": "https://developer.chrome.com/docs/lighthouse/overview",
        "discapacidad": "Visual, motora y cognitiva",
        "estado": "Recomendada para agregar (QA)",
    },
    {
        "nombre": "axe-core",
        "version": "v4.12.1 — Deque Systems",
        "url": "https://github.com/dequelabs/axe-core",
        "discapacidad": "Visual, motora, cognitiva y auditiva",
        "estado": "Recomendada para agregar (QA)",
    },
]

# Tabla
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.autofit = False

widths = [Cm(6.5), Cm(5.5), Cm(5.5)]
headers = [
    "Nombre y versión de la herramienta",
    "Enlace de la herramienta (fuente)",
    "Tipo de discapacidad orientada",
]

hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    set_cell_text(hdr_cells[i], h, bold=True, size=10)
    shade_cell(hdr_cells[i], "D9E1F2")
    hdr_cells[i].width = widths[i]

for herr in herramientas:
    row_cells = table.add_row().cells

    # Columna 1: nombre y versión
    p = row_cells[0].paragraphs[0]
    run = p.add_run(herr["nombre"])
    run.bold = True
    run.font.size = Pt(10)
    p2 = row_cells[0].add_paragraph()
    run2 = p2.add_run(herr["version"])
    run2.font.size = Pt(9)
    run2.italic = True

    # Columna 2: enlace
    row_cells[1].text = ""
    p_link = row_cells[1].paragraphs[0]
    add_hyperlink(p_link, herr["url"], herr["url"])
    for run_ in p_link.runs:
        run_.font.size = Pt(9)

    # Columna 3: discapacidad
    set_cell_text(row_cells[2], herr["discapacidad"], size=10)

    for c in row_cells:
        c.width = widths[0]

doc.add_paragraph()

# Nota final
nota_titulo = doc.add_paragraph()
run = nota_titulo.add_run("Nota:")
run.bold = True

nota = doc.add_paragraph(
    "Como ya se mencionó, UserWay está integrado en el proyecto y le da al "
    "usuario final el lector de pantalla, el ajuste de contraste, el cambio "
    "de tamaño de texto y la navegación por teclado. WAVE, Lighthouse y "
    "axe-core, en cambio, las proponemos como herramientas de QA para usar "
    "durante el desarrollo: sirven para revisar que las páginas cumplan con "
    "WCAG 2.2 antes de publicarlas. Como evalúan el HTML/CSS/JS ya "
    "renderizado en el navegador, funcionan sin importar que el backend esté "
    "hecho en PHP."
)

# ─────────────────────────────────────────────────────────
# Sección: Métricas de Software del Proyecto
# ─────────────────────────────────────────────────────────
doc.add_page_break()

heading_metricas = doc.add_heading(
    'Métricas de Software del Proyecto Tinkuy',
    level=2,
)

doc.add_paragraph(
    "Además de la accesibilidad, también evaluamos la calidad interna del "
    "código del proyecto. Para esto usamos cloc v1.64 sobre los 84 archivos "
    "PHP y un script propio que calcula el fan-in/fan-out de las funciones y "
    "la profundidad de anidamiento. A continuación se muestran los "
    "resultados."
)

doc.add_paragraph()

# --- Resumen general (equivalente CLOC) ---
doc.add_heading('Resumen General (equivalente CLOC)', level=3)

cloc_widths = [Cm(5), Cm(3), Cm(3), Cm(3)]
cloc_table = doc.add_table(rows=1, cols=4)
cloc_table.style = 'Table Grid'
cloc_table.autofit = False

hdr_cells = cloc_table.rows[0].cells
for i, h in enumerate(['Métrica', 'Cantidad', 'Porcentaje', 'Estado']):
    set_cell_text(hdr_cells[i], h, bold=True, size=10)
    shade_cell(hdr_cells[i], "D9E1F2")
    hdr_cells[i].width = cloc_widths[i]

cloc_rows = [
    ('Total líneas', '14,469', '100%', '—', None),
    ('Código fuente', '12,058', '83.3%', '⚠ ALTO', 'FCE4D6'),
    ('Comentarios', '946', '6.5%', '⚠ BAJO', 'FCE4D6'),
    ('Líneas en blanco', '1,465', '10.1%', '✓ NORMAL', 'E2EFDA'),
]
for metrica, cantidad, pct, estado, color in cloc_rows:
    row_cells = cloc_table.add_row().cells
    set_cell_text(row_cells[0], metrica, size=10)
    set_cell_text(row_cells[1], cantidad, size=10)
    set_cell_text(row_cells[2], pct, size=10)
    set_cell_text(row_cells[3], estado, size=10)
    if color:
        shade_cell(row_cells[3], color)
    for i, c in enumerate(row_cells):
        c.width = cloc_widths[i]

doc.add_paragraph(
    "El porcentaje de comentarios (6.5%) está por debajo del 20% "
    "recomendado, así que el proyecto queda como poco documentado. En "
    "cambio, el porcentaje de líneas en blanco (10.1%) sí está dentro del "
    "rango esperado (hasta 25%)."
)

doc.add_paragraph()

# --- Top 5 Fan-in ---
doc.add_heading('Top 5 Fan-in: funciones más reutilizadas', level=3)

fanin_widths = [Cm(1), Cm(4.5), Cm(2), Cm(8.5)]
fanin_table = doc.add_table(rows=1, cols=4)
fanin_table.style = 'Table Grid'
fanin_table.autofit = False

hdr_cells = fanin_table.rows[0].cells
for i, h in enumerate(['#', 'Función', 'Fan-in', 'Interpretación']):
    set_cell_text(hdr_cells[i], h, bold=True, size=10)
    shade_cell(hdr_cells[i], "D9E1F2")
    hdr_cells[i].width = fanin_widths[i]

fanin_rows = [
    ('1', 'tieneColumnaEstado', '3', 'Verifica si la tabla "mensajes" tiene columna de estado'),
    ('2', 'contarIntentos', '2', 'Control de intentos de inicio de sesión (seguridad)'),
    ('3', 'validarClave', '2', 'Validación de contraseña en registro y edición de usuario'),
    ('4', 'validarEmail', '2', 'Validación de correo en registro y edición de usuario'),
    ('5', 'validarNombre', '2', 'Validación de nombre/apellido en formularios'),
]
for num, fn, fi, interp in fanin_rows:
    row_cells = fanin_table.add_row().cells
    set_cell_text(row_cells[0], num, size=10)
    set_cell_text(row_cells[1], fn, size=10)
    set_cell_text(row_cells[2], fi, size=10)
    set_cell_text(row_cells[3], interp, size=9)
    for i, c in enumerate(row_cells):
        c.width = fanin_widths[i]

doc.add_paragraph(
    "Un fan-in alto indica que una función se usa desde varios puntos del "
    "sistema, lo que suele ser señal de buena cohesión. En Tinkuy los "
    "valores son bajos (máximo 3), algo razonable dado que el proyecto tiene "
    "poca lógica compartida entre módulos: cada controlador resuelve casi "
    "todo internamente."
)

doc.add_paragraph()

# --- Top 5 Fan-out ---
doc.add_heading('Top 5 Fan-out: funciones con más dependencias', level=3)

fanout_widths = [Cm(1), Cm(4.5), Cm(7), Cm(2.5)]
fanout_table = doc.add_table(rows=1, cols=4)
fanout_table.style = 'Table Grid'
fanout_table.autofit = False

hdr_cells = fanout_table.rows[0].cells
for i, h in enumerate(['#', 'Función', 'Archivo', 'Fan-out']):
    set_cell_text(hdr_cells[i], h, bold=True, size=10)
    shade_cell(hdr_cells[i], "D9E1F2")
    hdr_cells[i].width = fanout_widths[i]

fanout_rows = [
    ('1', 'crearUsuario', 'AdminUsuariosController.php', '5'),
    ('2', 'generar', 'ReportesController.php', '5'),
    ('3', 'procesarPago', 'PaymentController.php', '4'),
    ('4', 'procesarAcciones', 'MensajesController.php', '2'),
    ('5', 'send_mail', 'mailer_config.php', '2'),
]
for num, fn, archivo, fo in fanout_rows:
    row_cells = fanout_table.add_row().cells
    set_cell_text(row_cells[0], num, size=10)
    set_cell_text(row_cells[1], fn, size=10)
    set_cell_text(row_cells[2], archivo, size=10)
    set_cell_text(row_cells[3], fo, size=10)
    shade_cell(row_cells[3], "E2EFDA")
    for i, c in enumerate(row_cells):
        c.width = fanout_widths[i]

doc.add_paragraph(
    "El fan-out mide cuántas funciones distintas llama una función. Acá el "
    "máximo es 5 (crearUsuario y generar), un valor razonable que indica que "
    "ningún módulo concentra demasiadas dependencias internas."
)

doc.add_paragraph()

# --- Profundidad de anidamiento condicional ---
doc.add_heading('Profundidad de Anidamiento Condicional', level=3)

doc.add_paragraph("Límite máximo recomendado: 3 niveles.")

nest_widths = [Cm(9), Cm(3), Cm(4)]
nest_table = doc.add_table(rows=1, cols=3)
nest_table.style = 'Table Grid'
nest_table.autofit = False

hdr_cells = nest_table.rows[0].cells
for i, h in enumerate(['Archivo', 'Nivel máx.', 'Estado']):
    set_cell_text(hdr_cells[i], h, bold=True, size=10)
    shade_cell(hdr_cells[i], "D9E1F2")
    hdr_cells[i].width = nest_widths[i]

nest_rows = [
    ('AdminProductosController.php', '10'),
    ('VendedorController.php', '10'),
    ('AuthController.php', '6'),
    ('forgot_password.php', '6'),
    ('AdminController.php', '5'),
    ('AdminCuponesController.php', '5'),
    ('AdminUsuariosController.php', '5'),
    ('RegisterController.php', '5'),
    ('UserController.php', '5'),
    ('reset_password.php', '5'),
]
for archivo, nivel in nest_rows:
    row_cells = nest_table.add_row().cells
    set_cell_text(row_cells[0], archivo, size=10)
    set_cell_text(row_cells[1], nivel, size=10)
    set_cell_text(row_cells[2], 'REFACTORIZAR', size=10)
    shade_cell(row_cells[2], "FCE4D6")
    for i, c in enumerate(row_cells):
        c.width = nest_widths[i]

doc.add_paragraph(
    "17 de los 84 archivos (un 20% del proyecto) superan ese límite. Los "
    "casos más extremos son AdminProductosController.php y "
    "VendedorController.php, que llegan a 10 niveles; ambos serían los "
    "primeros candidatos a refactorizar para mejorar la legibilidad del "
    "código."
)

doc.add_paragraph()

# --- Estado general y conclusiones ---
doc.add_heading('Conclusiones', level=3)

estado_p = doc.add_paragraph()
estado_p.add_run('Estado general del proyecto: ')
run_estado = estado_p.add_run('REGULAR')
run_estado.bold = True

for bullet in [
    "El espaciado del código está bien: 10.1% de líneas en blanco, dentro "
    "del rango recomendado.",
    "Falta documentación: solo 6.5% de comentarios, cuando lo recomendado "
    "es al menos 10%.",
    "La densidad de código es alta (83.3%), por encima del 75% recomendado "
    "como máximo.",
]:
    doc.add_paragraph(bullet, style='List Bullet')

doc.add_paragraph()

doc.add_paragraph(
    "Tomando en cuenta el fan-in y el fan-out, ambos con valores bajos "
    "(máximo 3 y 5), el proyecto no presenta problemas serios de "
    "acoplamiento entre módulos. El punto más débil sigue siendo la "
    "documentación, junto con el anidamiento excesivo en "
    "VendedorController.php y AdminProductosController.php (10 niveles "
    "cada uno). Por eso, de cara a una posible mejora, recomendamos "
    "priorizar dos cosas: agregar más comentarios explicativos en el "
    "código, sobre todo en esos dos controladores, y simplificar sus "
    "bloques condicionales más anidados para hacerlos más legibles."
)

output_path = "Herramientas_Accesibilidad_Tinkuy.docx"
doc.save(output_path)
print(f"Documento generado: {output_path}")
